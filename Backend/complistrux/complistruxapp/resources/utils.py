from docx import Document
from docx.shared import Inches, Pt
import io
from pathlib import Path
from PIL import Image
import magic
import re
import zipfile
import pathlib
import os

def generateDocuments(data):

    image = str(data.image)
    image2 = re.sub(".*?/", "", image)
    


    imageurl = f"http://127.0.0.1:8000{data.image}"

    def image_to_jpg(image_path_or_stream):
        f = io.BytesIO()
        if isinstance(image_path_or_stream, str):
            path = Path(image_path_or_stream)
            if path.suffix in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
                f = open(image_path_or_stream, mode='rb')
            else:
                Image.open(image_path_or_stream).convert('RGB').save(f, format='JPEG')
        else:
            buffer = image_path_or_stream.read()
            mime_type = magic.from_buffer(buffer, mime=True)
            if mime_type in {'image/jpeg', 'image/png', 'image/gif', 'image/tiff', 'image/x-ms-bmp'}:
                f = image_path_or_stream
            else:
                Image.open(io.BytesIO(buffer)).convert('RGB').save(f, format='JPEG')
        return f
    # Build paths inside the project like this: BASE_DIR / 'subdir'.
    BASE_DIR = Path(__file__).resolve().parent.parent  
    BASE_DIR_IMAGE = Path(__file__).resolve().parent.parent.parent 
    
    # template_file_path = 'testdocforscript.docx'
    # output_file_path = 'result.docx'
    directory2 = BASE_DIR.joinpath("gendocstemps")
    
    delete_file_path = BASE_DIR.joinpath("gendocsoutput")

    for file in delete_file_path.iterdir():
        os.remove(file)

    for file_path in directory2.iterdir():

        path = pathlib.PurePath(file_path)
        print(path.name)

        template_file_path = file_path
        output_file_path = BASE_DIR.joinpath("gendocsoutput", f"{data.company_name} {path.name}")
        textvariables = {
            "[COMPANY NAME]": f"{data.company_name}",
            "[SIEM Solution]": f"{data.siem_solution}",
            "[End Point Manager]": "Microsoft Endpoint Manager",
            "[Firewall Solution]": f"{data.av_solution}",
            "[Access Control Solution]": f"{data.access_control_solution}",
        }

        template_document = Document(template_file_path)

        for variable_key, variable_value in textvariables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        header = template_document.sections[0].header
        paragraph = header.paragraphs[0]
        logo_run = paragraph.add_run()
        
        logo_run.add_picture(image_to_jpg(open(BASE_DIR_IMAGE.joinpath("media", "client_images", image2), mode='rb')) ,width=Inches(1.5))

        
        template_document.save(output_file_path)

        directory = BASE_DIR.joinpath("gendocsoutput")

        
        with zipfile.ZipFile("clientdocs.zip", mode="w") as archive:
            for file_path in directory.iterdir():
                archive.write(file_path, arcname=file_path.name)



def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)



# if __name__ == '__main__':
#     main()